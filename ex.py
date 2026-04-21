#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
分红审核示例脚本

流程:
1. 浓缩文档内容提取分红信息
2. 输入LLM进行分红审核
3. 输出审核json结果
4. 写入数据库
5. 触发kafka

使用方式:
    python ex.py --doc_id 101 --env dev
    python ex.py --docx_path /path/to/file.docx
    python ex.py --docx_url http://xxx/file.docx
"""

import os
import sys
import argparse
import json
import tempfile
import urllib.request
from pathlib import Path

# 添加项目根目录到路径
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

try:
    from docx import Document
except ImportError:
    print("Error: python-docx is required. Install with: pip install python-docx", file=sys.stderr)
    sys.exit(1)

from app.config import cfg
from app.logger import logger
from modules.common.db_client import exec_query_df
from modules.common.encryption import decrypt_text


# ============================================================
# 关键词定义 (参考 adp_py/ex.py)
# ============================================================
KEYWORDS = {
    "dividend": [
        "分红", "利润分配", "股息", "股利", "现金分红", "分红率",
        "每10股", "派息", "送红股", "转增", "分配预案", "分配方案",
        "每10股派", "每10股转", "现金利润分配", "未分配利润"
    ],
    "profit": [
        "净利润", "归母净利润", "归属于上市公司股东", "扣非净利润",
        "扣除非经常性损益", "经常性损益", "基本每股收益", "稀释每股收益",
        "加权平均净资产收益率", "扣非后基本每股收益"
    ],
    "revenue": [
        "营业收入", "营业成本", "主营业务收入", "其他业务收入"
    ],
    "asset": [
        "总资产", "净资产", "归属于上市公司股东的净资产", "总股本",
        "股本", "注册资本"
    ],
    "cashflow": [
        "经营活动现金流", "投资活动现金流", "筹资活动现金流",
        "现金及现金等价物"
    ],
    "company": [
        "公司名称", "公司简称", "股票代码", "年度报告", "公司代码"
    ]
}


# ============================================================
# 文档浓缩模块
# ============================================================
def safe_iter_table(table):
    """安全地遍历表格，处理损坏的表格"""
    try:
        for row in table.rows:
            try:
                yield [cell.text.strip() for cell in row.cells]
            except (ValueError, RuntimeError):
                break
    except (ValueError, RuntimeError):
        return


def extract_relevant_content_from_doc(doc: Document) -> str:
    """从DOCX文档中提取与分红、净利润相关的财务信息"""
    lines = []
    found_sections = set()

    # 1. 提取段落中的相关信息
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        for category, kws in KEYWORDS.items():
            for kw in kws:
                if kw in text:
                    if text not in found_sections:
                        lines.append(f"[段落] {text}")
                        found_sections.add(text)
                    break

    # 2. 提取表格中的相关信息
    for t_idx, table in enumerate(doc.tables):
        rows = list(safe_iter_table(table))
        if not rows:
            continue

        relevant_rows = []
        for row in rows:
            row_text = " | ".join(row)
            for category, kws in KEYWORDS.items():
                for kw in kws:
                    if kw in row_text:
                        relevant_rows.append(" | ".join(cell for cell in row if cell))
                        break

        if relevant_rows:
            lines.append(f"\n[表格 {t_idx}]")
            lines.extend(relevant_rows)

    return "\n".join(lines)


def extract_relevant_content_from_chunks(chunks: list) -> str:
    """从数据库chunk中提取与分红、净利润相关的财务信息"""
    lines = []
    found_sections = set()

    for chunk in chunks:
        text = chunk.get('text', '')
        if not text:
            continue

        # 检查是否包含关键词
        is_relevant = False
        for category, kws in KEYWORDS.items():
            for kw in kws:
                if kw in text:
                    is_relevant = True
                    break
            if is_relevant:
                break

        if is_relevant and text not in found_sections:
            lines.append(f"[chunk_{chunk.get('id')}] {text}")
            found_sections.add(text)

    return "\n".join(lines)


# ============================================================
# 数据库操作模块
# ============================================================
def get_chunks_from_db(doc_id: int, encrypt_flag: bool = True) -> list:
    """从数据库获取文档chunk"""
    sql = """
        SELECT id, paraid, text, metadata
        FROM analysis_doc_chunk_info_ib
        WHERE doc_id = {} AND audit_type = 2
        ORDER BY id
    """.format(doc_id)

    try:
        df = exec_query_df(sql)
        if df.empty:
            return []

        chunks = []
        for _, row in df.iterrows():
            text = row['text']
            if encrypt_flag:
                try:
                    text = decrypt_text(text)
                except:
                    pass

            chunks.append({
                'id': row['id'],
                'paraid': row['paraid'],
                'text': text,
                'metadata': row['metadata']
            })

        return chunks
    except Exception as e:
        print("查询chunk失败: {}".format(e))
        return []


def get_doc_url_from_db(doc_id: int) -> str:
    """从数据库获取文档URL"""
    sql = """
        SELECT url FROM analysis_doc_chunk_info_ib
        WHERE doc_id = {} LIMIT 1
    """.format(doc_id)

    try:
        df = exec_query_df(sql)
        if df.empty:
            return None
        return df.values[0][0]
    except Exception as e:
        print("查询URL失败: {}".format(e))
        return None


# ============================================================
# 文件获取模块
# ============================================================
def get_docx_path(docx_path: str) -> tuple:
    """获取 docx 文件路径，支持本地文件和 URL"""
    path = Path(docx_path)

    # 本地文件
    if path.exists():
        return path, False

    # URL 下载
    if docx_path.startswith(('http://', 'https://', 'ftp://')):
        try:
            with urllib.request.urlopen(docx_path, timeout=30) as response:
                content = response.read()
        except Exception as e:
            print(f"Error: Failed to download from URL: {e}", file=sys.stderr)
            return None, False

        temp_dir = Path(tempfile.gettempdir())
        temp_path = temp_dir / f"dividend_docx_{id(docx_path)}.docx"
        temp_path.write_bytes(content)
        return temp_path, True

    print(f"Error: File not found: {docx_path}", file=sys.stderr)
    return None, False


# ============================================================
# LLM审核模块
# ============================================================
def llm_audit(condensed_text: str, doc_id: int) -> list:
    """调用LLM进行分红审核"""
    from modules.dividend_audit.llm_auditor import DividendAuditor

    auditor = DividendAuditor()

    # 构建chunk格式
    chunks = [{'id': 0, 'text': condensed_text}]

    return auditor.audit(chunks, doc_id)


# ============================================================
# 数据库写入模块
# ============================================================
def save_to_db(doc_id: int, audit_result: list):
    """将审核结果写入数据库"""
    from modules.dividend_audit.db_writer import DividendDBWriter

    writer = DividendDBWriter()
    writer.save(doc_id, audit_result)
    print("审核结果已写入数据库, doc_id: {}".format(doc_id))


# ============================================================
# Kafka通知模块
# ============================================================
def send_to_kafka(doc_id: int, audit_result: list):
    """发送审核结果到Kafka"""
    try:
        from kafka import KafkaProducer
        import json

        kafka_config = cfg.KafkaConfig
        producer = KafkaProducer(
            bootstrap_servers=kafka_config['bootstrap_servers'],
            value_serializer=lambda v: json.dumps(v, ensure_ascii=False).encode('utf-8')
        )

        topic = kafka_config.get('result_topic', 'jz_ib_analysis_document_audit_topic_v2_push')

        message = {
            'doc_id': doc_id,
            'audit_type': 'dividend',
            'status': 200,
            'result_count': len(audit_result),
            'result': audit_result
        }

        producer.send(topic, value=message, key=str(doc_id).encode('utf-8'))
        producer.flush()
        producer.close()

        print("审核结果已发送到Kafka, topic: {}, doc_id: {}".format(topic, doc_id))
    except Exception as e:
        print("Kafka发送失败: {}".format(e))


# ============================================================
# 主流程
# ============================================================
def process_by_doc_id(doc_id: int, encrypt_flag: bool = True):
    """通过doc_id处理: 从数据库读取chunk"""
    print("\n" + "=" * 60)
    print("Step 1: 从数据库获取chunk")
    print("=" * 60)

    chunks = get_chunks_from_db(doc_id, encrypt_flag)
    if not chunks:
        print("未找到chunk数据, doc_id: {}".format(doc_id))
        return None

    print("获取chunk数: {}".format(len(chunks)))

    print("\n" + "=" * 60)
    print("Step 2: 浓缩文档内容提取分红信息")
    print("=" * 60)

    condensed_text = extract_relevant_content_from_chunks(chunks)
    if not condensed_text:
        print("未找到分红相关内容")
        return []

    print("浓缩后文本长度: {} 字符".format(len(condensed_text)))
    print("\n浓缩内容预览:")
    print("-" * 40)
    print(condensed_text[:500] + "..." if len(condensed_text) > 500 else condensed_text)

    print("\n" + "=" * 60)
    print("Step 3: 输入LLM进行分红审核")
    print("=" * 60)

    audit_result = llm_audit(condensed_text, doc_id)
    print("审核完成, 发现问题数: {}".format(len(audit_result)))

    for i, item in enumerate(audit_result[:3]):
        print("\n问题 {}:".format(i + 1))
        print("  - 类型: {}".format(item.get('type')))
        print("  - 内容: {}".format(item.get('content', '')[:100]))

    print("\n" + "=" * 60)
    print("Step 4: 写入数据库")
    print("=" * 60)

    save_to_db(doc_id, audit_result)

    print("\n" + "=" * 60)
    print("Step 5: 触发Kafka")
    print("=" * 60)

    send_to_kafka(doc_id, audit_result)

    return audit_result


def process_by_docx_path(docx_path: str):
    """通过docx文件路径处理"""
    print("\n" + "=" * 60)
    print("Step 1: 获取文档")
    print("=" * 60)

    path, is_temp = get_docx_path(docx_path)
    if not path:
        return None

    print("文档路径: {}".format(path))

    print("\n" + "=" * 60)
    print("Step 2: 解析DOCX并浓缩")
    print("=" * 60)

    try:
        doc = Document(path)
        condensed_text = extract_relevant_content_from_doc(doc)
    except Exception as e:
        print("解析DOCX失败: {}".format(e))
        return None
    finally:
        if is_temp:
            path.unlink(missing_ok=True)

    if not condensed_text:
        print("未找到分红相关内容")
        return []

    print("浓缩后文本长度: {} 字符".format(len(condensed_text)))
    print("\n浓缩内容预览:")
    print("-" * 40)
    print(condensed_text[:500] + "..." if len(condensed_text) > 500 else condensed_text)

    print("\n" + "=" * 60)
    print("Step 3: 输入LLM进行分红审核")
    print("=" * 60)

    audit_result = llm_audit(condensed_text, 0)
    print("审核完成, 发现问题数: {}".format(len(audit_result)))

    return audit_result


# ============================================================
# 命令行入口
# ============================================================
def main():
    parser = argparse.ArgumentParser(description='分红审核示例脚本')
    parser.add_argument('--doc_id', type=int, help='文档ID (从数据库读取)')
    parser.add_argument('--docx_path', type=str, help='DOCX文件路径')
    parser.add_argument('--docx_url', type=str, help='DOCX文件URL')
    parser.add_argument('--env', type=str, default='dev', help='环境 (dev/qa/uat/prd)')
    parser.add_argument('--encrypt', type=bool, default=True, help='是否解密')

    args = parser.parse_args()

    # 设置环境
    os.environ['ENV'] = args.env

    print("=" * 60)
    print("分红审核流程")
    print("=" * 60)

    result = None

    if args.doc_id:
        result = process_by_doc_id(args.doc_id, args.encrypt)
    elif args.docx_path:
        result = process_by_docx_path(args.docx_path)
    elif args.docx_url:
        result = process_by_docx_path(args.docx_url)
    else:
        print("请指定 --doc_id 或 --docx_path 或 --docx_url")
        parser.print_help()
        sys.exit(1)

    print("\n" + "=" * 60)
    print("审核完成")
    print("=" * 60)

    if result is not None:
        print("\n审核结果JSON:")
        print(json.dumps(result, ensure_ascii=False, indent=2))

    return result


if __name__ == '__main__':
    main()
