from docx import Document
from io import BytesIO
import requests
from app.logger import logger
from app.config import cfg
from modules.common.s3_client import s3_client
from modules.common.db_client import exec_query_df
from .text_condenser import TextCondenser
from .llm_auditor import DividendAuditor
from .db_writer import DividendDBWriter

class DividendProcessor:
    """分红审核处理器"""

    def __init__(self):
        self.condenser = TextCondenser()
        self.auditor = DividendAuditor()
        self.db_writer = DividendDBWriter()

    def process(self, args):
        """
        主处理流程

        Args:
            args: 请求参数
                - doc_id: 文档ID
                - encrypt_flag: 是否加密
                - callback_url: 回调URL

        Returns:
            list: 审核结果列表
        """
        doc_id = args.get("doc_id")
        encrypt_flag = args.get("encrypt_flag", True)
        callback_url = args.get("callback_url")

        logger.info("开始处理 doc_id: {}".format(doc_id))

        # Step 1: 获取文档URL
        doc_url = self._get_doc_url(doc_id)
        if not doc_url:
            return self._error_result("文档不存在", doc_id)

        # Step 2: 获取文档内容
        try:
            doc_content = s3_client.get_file_content(doc_url, encrypt_flag)
        except Exception as e:
            return self._error_result("获取文档内容失败: {}".format(str(e)), doc_id)

        # Step 3: 解析DOCX
        paragraphs = self._parse_docx(doc_content)
        if not paragraphs:
            return self._error_result("文档解析失败或内容为空", doc_id)

        logger.info("doc_id: {} 解析出段落数: {}".format(doc_id, len(paragraphs)))

        # Step 4: 文本浓缩
        condensed_text = self.condenser.condense(paragraphs)
        if not condensed_text:
            return self._error_result("未找到分红相关内容", doc_id)

        logger.info("doc_id: {}, 浓缩后长度: {}".format(doc_id, len(condensed_text)))

        # Step 5: LLM审核
        audit_result = self.auditor.audit(condensed_text, doc_id)

        # Step 6: 结果入库
        try:
            self.db_writer.save(doc_id, audit_result)
        except Exception as e:
            logger.error("结果入库失败: {}".format(e))

        # Step 7: 回调通知
        if callback_url:
            self._callback(callback_url, doc_id, audit_result)

        logger.info("doc_id: {} 处理完成, 发现问题数: {}".format(doc_id, len(audit_result)))

        return audit_result

    def _get_doc_url(self, doc_id):
        """从数据库获取文档URL"""
        sql = """
            SELECT url FROM analysis_doc_chunk_info_ib
            WHERE doc_id = {} LIMIT 1
        """.format(doc_id)

        try:
            df = exec_query_df(sql)
            if df.empty:
                logger.error("doc_id: {} 未找到文档URL".format(doc_id))
                return None
            return df.values[0][0]
        except Exception as e:
            logger.error("查询文档URL失败: {}".format(e))
            return None

    def _parse_docx(self, doc_content):
        """解析DOCX文件"""
        try:
            doc = Document(doc_content)
        except Exception as e:
            logger.error("打开DOCX文件失败: {}".format(e))
            return []

        paragraphs = []

        # 提取段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append({
                    'text': text,
                    'style': para.style.name if para.style else None
                })

        # 提取表格
        for table in doc.tables:
            table_text = self._extract_table(table)
            if table_text:
                paragraphs.append({
                    'text': table_text,
                    'style': 'table'
                })

        return paragraphs

    def _extract_table(self, table):
        """提取表格内容"""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        return "\n".join(rows)

    def _callback(self, url, doc_id, result):
        """回调通知"""
        try:
            response = requests.post(
                url,
                json={
                    'doc_id': doc_id,
                    'status': 'completed',
                    'result_count': len(result),
                    'has_error': len(result) > 0
                },
                timeout=10
            )
            logger.info("回调成功: {}".format(url))
        except Exception as e:
            logger.warning("回调失败: {}".format(e))

    def _error_result(self, message, doc_id=None):
        """返回错误结果"""
        logger.error("doc_id: {}, 错误: {}".format(doc_id, message))
        return [{
            'type': '系统错误',
            'content': message,
            'reason': message,
            'suggestion': '请检查文档或联系管理员',
            'l1_type': '系统错误',
            'l2_type': '分红数据审核'
        }]
