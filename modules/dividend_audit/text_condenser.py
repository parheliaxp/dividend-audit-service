import re
from app.logger import logger
from app.config import cfg

try:
    from docx import Document
except ImportError:
    Document = None


class TextCondenser:
    """文本浓缩器 - 提取分红相关内容"""

    # 综合关键词定义 (参考 adp_py/ex.py)
    KEYWORDS = {
        "dividend": [
            "分红", "利润分配", "股息", "股利", "现金分红", "分红率",
            "每10股", "派息", "送红股", "转增", "分配预案", "分配方案",
            "每10股派", "每10股转", "现金利润分配", "未分配利润"
        ],
        "profit": [
            "净利润", "归母净利润", "归属于上市公司股东", "扣非净利润",
            "扣除非经常性损益", "经常性损益", "基本每股收益", "稀释每股收益",
            "加权平均净资产收益率", "扣非后基本每股收益"
        ],
        "revenue": [
            "营业收入", "营业成本", "主营业务收入", "其他业务收入"
        ],
        "asset": [
            "总资产", "净资产", "归属于上市公司股东的净资产", "总股本",
            "股本", "注册资本"
        ],
        "cashflow": [
            "经营活动现金流", "投资活动现金流", "筹资活动现金流",
            "现金及现金等价物"
        ],
        "company": [
            "公司名称", "公司简称", "股票代码", "年度报告", "公司代码"
        ]
    }

    # 需要排除的无关内容
    EXCLUDE_PATTERNS = [
        r'本报告.*仅供参考',
        r'投资者.*注意.*风险',
        r'本公司.*承诺',
        r'不构成.*投资建议',
    ]

    def __init__(self, max_length=None):
        self.max_length = max_length or cfg.DividendAuditConfig['max_text_length']

    def filter_dividend_chunks(self, chunks):
        """
        筛选分红相关的chunk

        Args:
            chunks: chunk列表 [{'id': int, 'paraid': str, 'text': str, 'metadata': str}, ...]

        Returns:
            list: 分红相关的chunk
        """
        dividend_chunks = []

        for chunk in chunks:
            text = chunk.get('text', '')

            # 检查是否包含分红关键词
            if self._contains_dividend_keywords(text):
                dividend_chunks.append(chunk)
                continue

            # 检查是否包含财务数据模式
            if self._contains_financial_data(text):
                dividend_chunks.append(chunk)

        logger.info("筛选分红chunk: {} / {}".format(len(dividend_chunks), len(chunks)))
        return dividend_chunks

    def _contains_dividend_keywords(self, text):
        """检查文本是否包含分红关键词"""
        for category, kws in self.KEYWORDS.items():
            for kw in kws:
                if kw in text:
                    return True
        return False

    def _contains_financial_data(self, text):
        """检查是否包含财务数据模式"""
        patterns = [
            r'\d+\.?\d*\s*[万亿]?元',      # 金额: 100万元, 1.5亿元
            r'每股\s*\d+\.?\d*\s*元?',      # 每股: 每股0.5元
            r'\d+\.?\d*%',                  # 百分比: 50%
            r'\d{4}年\d{1,2}月\d{1,2}日',   # 日期: 2024年1月15日
        ]

        for pattern in patterns:
            if re.search(pattern, text):
                return True
        return False

    def condense(self, chunks):
        """
        浓缩文本，提取分红相关内容

        Args:
            chunks: chunk列表

        Returns:
            str: 浓缩后的文本
        """
        # Step 1: 筛选分红相关chunk
        dividend_chunks = self.filter_dividend_chunks(chunks)

        # Step 2: 提取文本
        texts = []
        for chunk in dividend_chunks:
            text = chunk.get('text', '')
            # 清洗无关内容
            text = self._clean_content(text)
            if text:
                texts.append(text)

        # Step 3: 合并并截断
        condensed_text = self._merge_and_truncate(texts)

        return condensed_text

    def _clean_content(self, text):
        """清洗无关内容"""
        # 移除匹配排除模式的内容
        for pattern in self.EXCLUDE_PATTERNS:
            text = re.sub(pattern, '', text)

        text = text.strip()
        return text if len(text) > 5 else ''

    def _merge_and_truncate(self, texts):
        """合并文本并截断到最大长度"""
        merged = "\n\n".join(texts)

        if len(merged) > self.max_length:
            merged = merged[:self.max_length]
            logger.warning("文本已截断到 {} 字符".format(self.max_length))

        return merged

    def get_chunk_ids(self, chunks):
        """获取chunk ID列表"""
        return [chunk.get('id') for chunk in chunks]

    # ============================================================
    # DOCX文档解析方法 (参考 adp_py/ex.py)
    # ============================================================
    @staticmethod
    def safe_iter_table(table):
        """安全地遍历表格，处理损坏的表格"""
        try:
            for row in table.rows:
                try:
                    yield [cell.text.strip() for cell in row.cells]
                except (ValueError, RuntimeError):
                    break
        except (ValueError, RuntimeError):
            return

    def extract_relevant_content_from_doc(self, doc) -> str:
        """从DOCX文档中提取与分红、净利润相关的财务信息"""
        lines = []
        found_sections = set()

        # 1. 提取段落中的相关信息
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            for category, kws in self.KEYWORDS.items():
                for kw in kws:
                    if kw in text:
                        if text not in found_sections:
                            lines.append(f"[段落] {text}")
                            found_sections.add(text)
                        break

        # 2. 提取表格中的相关信息
        for t_idx, table in enumerate(doc.tables):
            rows = list(self.safe_iter_table(table))
            if not rows:
                continue

            relevant_rows = []
            for row in rows:
                row_text = " | ".join(row)
                for category, kws in self.KEYWORDS.items():
                    for kw in kws:
                        if kw in row_text:
                            relevant_rows.append(" | ".join(cell for cell in row if cell))
                            break

            if relevant_rows:
                lines.append(f"\n[表格 {t_idx}]")
                lines.extend(relevant_rows)

        return "\n".join(lines)

    def extract_relevant_content_from_chunks(self, chunks: list) -> str:
        """从数据库chunk中提取与分红、净利润相关的财务信息"""
        lines = []
        found_sections = set()

        for chunk in chunks:
            text = chunk.get('text', '')
            if not text:
                continue

            # 检查是否包含关键词
            is_relevant = False
            for category, kws in self.KEYWORDS.items():
                for kw in kws:
                    if kw in text:
                        is_relevant = True
                        break
                if is_relevant:
                    break

            if is_relevant and text not in found_sections:
                lines.append(f"[chunk_{chunk.get('id')}] {text}")
                found_sections.add(text)

        return "\n".join(lines)

    def condense_from_docx(self, docx_path: str) -> str:
        """
        从DOCX文件路径提取分红相关内容

        Args:
            docx_path: DOCX文件路径

        Returns:
            str: 浓缩后的文本
        """
        if Document is None:
            raise ImportError("python-docx is required. Install with: pip install python-docx")

        from pathlib import Path
        path = Path(docx_path)

        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {docx_path}")

        doc = Document(path)
        condensed_text = self.extract_relevant_content_from_doc(doc)

        if len(condensed_text) > self.max_length:
            condensed_text = condensed_text[:self.max_length]
            logger.warning("文本已截断到 {} 字符".format(self.max_length))

        return condensed_text
